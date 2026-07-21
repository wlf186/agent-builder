"""Sandbox payload for the curated research-documents dependency bundle."""

from __future__ import annotations

from html.parser import HTMLParser
import hashlib
import json
import os
from pathlib import Path
import re
from typing import Iterable


MAX_DOCUMENT_BYTES = 16 * 1024 * 1024
MAX_PDF_PAGES = 512
MAX_DOCX_BLOCKS = 20_000
MAX_OFFSET_CHARS = 1_000_000
MAX_RETURN_CHARS = 4_096
_STAGED_NAME = re.compile(r"^research-input-[a-f0-9]{32}\.bin$")
_SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx", ".txt", ".md", ".html", ".htm"})


class _HTMLText(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._ignored = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs
        if tag in {"script", "style", "noscript"}:
            self._ignored += 1
        elif tag in {"p", "div", "section", "article", "br", "li", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"}:
            self._ignored = max(0, self._ignored - 1)
        elif tag in {"p", "div", "section", "article", "li"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._ignored == 0:
            self.parts.append(data)


def _clean(value: str) -> str:
    value = value.replace("\x00", " ")
    value = "".join(
        character
        for character in value
        if ord(character) >= 32 or character in "\n\t\r"
    )
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _slice_segments(
    segments: Iterable[str], offset: int, maximum: int
) -> tuple[str, int, bool]:
    selected: list[str] = []
    consumed = 0
    end = offset + maximum
    stopped_early = False
    for segment in segments:
        cleaned = _clean(segment)
        if not cleaned:
            continue
        value = cleaned + "\n"
        next_consumed = consumed + len(value)
        if next_consumed > offset and consumed < end:
            start_index = max(0, offset - consumed)
            end_index = min(len(value), end - consumed)
            selected.append(value[start_index:end_index])
        consumed = next_consumed
        if consumed >= end:
            stopped_early = True
            break
    return _clean("".join(selected)), consumed, stopped_early


def _pdf_segments(path: Path) -> tuple[Iterable[str], dict[str, object]]:
    from pypdf import PdfReader

    reader = PdfReader(path, strict=True)
    if reader.is_encrypted:
        raise ValueError("encrypted PDF is not supported")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError("PDF page count exceeds its limit")

    def values() -> Iterable[str]:
        for page in reader.pages:
            yield page.extract_text() or ""

    return values(), {"parser": "pypdf", "pages": len(reader.pages)}


def _docx_segments(path: Path) -> tuple[Iterable[str], dict[str, object]]:
    from docx import Document

    document = Document(path)

    def values() -> Iterable[str]:
        count = 0
        for paragraph in document.paragraphs:
            count += 1
            if count > MAX_DOCX_BLOCKS:
                raise ValueError("DOCX block count exceeds its limit")
            yield paragraph.text
        for table in document.tables:
            for row in table.rows:
                count += 1
                if count > MAX_DOCX_BLOCKS:
                    raise ValueError("DOCX block count exceeds its limit")
                yield "\t".join(cell.text for cell in row.cells)

    return values(), {"parser": "python-docx"}


def _text_segments(path: Path, suffix: str) -> tuple[Iterable[str], dict[str, object]]:
    raw = path.read_bytes()
    value = raw.decode("utf-8")
    if suffix in {".html", ".htm"}:
        parser = _HTMLText()
        parser.feed(value)
        parser.close()
        return parser.parts, {"parser": "html.parser"}
    return (value,), {"parser": "utf-8"}


def extract(arguments: object) -> dict[str, object]:
    if not isinstance(arguments, dict) or set(arguments) != {
        "original_path",
        "staged_name",
        "offset_chars",
        "max_chars",
        "content_digest",
    }:
        raise ValueError("research input is invalid")
    original_path = arguments.get("original_path")
    staged_name = arguments.get("staged_name")
    offset = arguments.get("offset_chars")
    maximum = arguments.get("max_chars")
    expected_digest = arguments.get("content_digest")
    if (
        not isinstance(original_path, str)
        or not original_path
        or not isinstance(staged_name, str)
        or _STAGED_NAME.fullmatch(staged_name) is None
        or not isinstance(offset, int)
        or isinstance(offset, bool)
        or not 0 <= offset <= MAX_OFFSET_CHARS
        or not isinstance(maximum, int)
        or isinstance(maximum, bool)
        or not 1 <= maximum <= MAX_RETURN_CHARS
        or not isinstance(expected_digest, str)
        or re.fullmatch(r"[a-f0-9]{64}", expected_digest) is None
    ):
        raise ValueError("research input is invalid")
    suffix = Path(original_path).suffix.lower()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise ValueError("document type is not supported")
    path = Path(staged_name)
    metadata = path.stat(follow_symlinks=False)
    if not path.is_file() or not 1 <= metadata.st_size <= MAX_DOCUMENT_BYTES:
        raise ValueError("staged document is invalid")
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    if digest != expected_digest:
        raise ValueError("staged document identity changed")
    if suffix == ".pdf":
        segments, metadata_value = _pdf_segments(path)
    elif suffix == ".docx":
        segments, metadata_value = _docx_segments(path)
    else:
        segments, metadata_value = _text_segments(path, suffix)
    content, observed, stopped_early = _slice_segments(segments, offset, maximum)
    returned = len(content)
    return {
        "schema_version": 1,
        "kind": "document_text",
        "path": original_path,
        "content_digest": digest,
        "content": content,
        "range": {
            "offset_chars": offset,
            "returned_chars": returned,
            "next_offset_chars": offset + returned,
        },
        "truncated": stopped_early or offset > 0 or observed > offset + returned,
        **metadata_value,
    }


def main() -> int:
    raw = os.environ.get("AGENT_BUILDER_SKILL_INPUT", "")
    if not raw or len(raw.encode("utf-8")) > 4_096:
        raise ValueError("research input is unavailable")
    result = extract(json.loads(raw))
    print(
        json.dumps(
            result,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ),
        flush=True,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
